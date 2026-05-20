# docx_utils.py - 实验报告 Word 模板填充工具
# 用法：
#   1. 复制本脚本到临时目录
#   2. 修改 main() 中的数据部分
#   3. 运行：.\.trae\uv tool run --from python-docx python docx_utils.py
# 依赖：python-docx

import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_cell(cell, text, font_name='Times New Roman', font_size=Pt(12),
             bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """设置表格单元格内容与格式"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_cell_styles(cell, font_name='Times New Roman', font_size=Pt(12),
                    bold=False):
    """仅修改单元格字体样式，不改变内容"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def fill_table(doc, table_index, data, start_row=0, start_col=0,
               font_name='Times New Roman', font_size=Pt(12),
               header_bold=True):
    """批量填充表格数据

    参数:
        doc: Document 对象
        table_index: 表格索引
        data: list[list] 二维数据
        start_row: 起始行
        start_col: 起始列
        font_name: 字体
        font_size: 字号
        header_bold: 表头是否加粗（第一行）
    """
    table = doc.tables[table_index]
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            row_idx = start_row + i
            col_idx = start_col + j
            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                is_header = (i == 0 and header_bold)
                set_cell(cell, cell_text, font_name=font_name,
                         font_size=font_size, bold=is_header)


def set_code_font(doc, start_idx, end_idx):
    """将指定段落范围的字体设为 Consolas 11pt"""
    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(11)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')


def append_text(paragraph, text, font_name='Times New Roman', font_size=Pt(12),
                bold=False, italic=False, color=None):
    """在已有段落末尾追加文本"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def insert_text(doc, para_index, text, font_name='Times New Roman',
                font_size=Pt(12), bold=False, align=None):
    """在指定段落位置写入文本（清空原内容后写入）"""
    p = doc.paragraphs[para_index]
    p.text = ''
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_picture_to_paragraph(doc, para_index, image_path, width=Cm(12)):
    """在指定段落插入图片"""
    p = doc.paragraphs[para_index]
    run = p.add_run()
    run.add_picture(image_path, width=width)


def find_paragraph_by_text(doc, text, contains=False):
    """查找包含指定文本的段落索引

    参数:
        doc: Document 对象
        text: 要搜索的文本
        contains: True=模糊匹配, False=精确匹配
    返回:
        段落索引, 或 None
    """
    for i, p in enumerate(doc.paragraphs):
        if contains:
            if text in p.text:
                return i
        else:
            if p.text.strip() == text:
                return i
    return None


def find_table_by_header(doc, header_text, col=0):
    """查找表头包含指定文本的表格索引

    参数:
        doc: Document 对象
        header_text: 表头文本
        col: 表头所在列
    返回:
        表格索引, 或 None
    """
    for ti, table in enumerate(doc.tables):
        if len(table.rows) > 0 and len(table.rows[0].cells) > col:
            if header_text in table.rows[0].cells[col].text:
                return ti
    return None


def docx_to_pdf(docx_path, pdf_path):
    """将 docx 转换为 PDF（需要安装 pywin32）"""
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()
        print(f"PDF 已保存: {pdf_path}")
    except ImportError:
        print("请先安装 pywin32: pip install pywin32")
        raise
    except Exception as e:
        print(f"PDF 转换失败: {e}")
        raise


def validate_tables(doc, filename="输出文档"):
    """验证文档表格结构"""
    print(f"=== {filename} 表格结构验证 ===")
    for ti, table in enumerate(doc.tables):
        print(f"  Table {ti}: {len(table.rows)}行 x {len(table.columns)}列")
        for ri, row in enumerate(table.rows):
            cells_text = [cell.text[:20] for cell in row.cells]
            print(f"    行{ri}: {cells_text}")


def validate_paragraphs(doc, filename="输出文档"):
    """验证文档段落"""
    print(f"=== {filename} 段落验证 ===")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"  P[{i}]: {p.text[:60]}...")


def backup_template(docx_path):
    """备份模板文件"""
    import shutil
    import os
    base, ext = os.path.splitext(docx_path)
    backup_path = f"{base}_备份{ext}"
    if not os.path.exists(backup_path):
        shutil.copy2(docx_path, backup_path)
        print(f"备份已创建: {backup_path}")
    return backup_path


if __name__ == '__main__':
    # ==========================================
    # 示例：第七次实验模板填充
    # 每次实验修改此部分即可
    # ==========================================

    template_path = r"F:\Project\机器学习\Machine Learning\Exp_7\report\机器学习_第七次实验报告模板.docx"
    student_name = "何欣航"
    student_id = "2023112560"
    exp_num = 7

    doc = docx.Document(template_path)

    # --- 1. 学生信息表 (Table 0) ---
    fill_table(doc, 0, [
        [f"姓名：{student_name}", f"学号：{student_id}", ""],
    ])

    # --- 2. 先验概率表 (Table 1) ---
    fill_table(doc, 1, [
        ["类别", "先验概率"],
        ["P(Y=1)", "0.40"],
        ["P(Y=2)", "0.40"],
        ["P(Y=3)", "0.20"],
    ])

    # --- 3. 高斯分布参数表 (Table 2) ---
    fill_table(doc, 2, [
        ["类别", "SepalLength\n均值", "SepalLength\n标准差",
         "SepalWidth\n均值", "SepalWidth\n标准差",
         "PetalLength\n均值", "PetalLength\n标准差",
         "PetalWidth\n均值", "PetalWidth\n标准差"],
        ["P(X|Y=1)", "5.04", "0.36", "3.44", "0.38",
         "1.46", "0.18", "0.23", "0.10"],
        ["P(X|Y=2)", "6.02", "0.52", "2.79", "0.33",
         "4.32", "0.45", "1.35", "0.21"],
        ["P(X|Y=3)", "6.56", "0.73", "2.92", "0.39",
         "5.66", "0.64", "2.04", "0.27"],
    ])

    # --- 4. 测试集准确率 (P[8]) ---
    insert_text(doc, 8,
                "测试集上准确率（保留两位小数）：0.92 (92.00%)",
                font_size=Pt(12))

    # --- 5. 代码字体设置 (P[12]~P[71]) ---
    set_code_font(doc, 12, 72)

    # --- 6. 追加实验原理文本到 P[73] ---
    append_text(doc.paragraphs[73],
                "（追加内容：本次实验详细分析了各特征的高斯分布参数...）")

    # --- 7. 保存 ---
    output_path = template_path
    doc.save(output_path)
    print(f"模板填充完成: {output_path}")

    # --- 8. 验证 ---
    validate_tables(docx.Document(output_path), output_path)
